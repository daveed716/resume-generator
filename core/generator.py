"""
Refactored generation logic extracted from generate_application.py.

All writer functions are copied verbatim. API calls accept explicit keys
and an on_status callback for GUI integration.
"""

import json
import os
import re
import textwrap
from typing import Callable, Optional

from anthropic import Anthropic
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
from fpdf import FPDF, XPos, YPos

from core.file_naming import build_base_paths, resolve_filename
from core.models import GeneratedFile, GenerationOptions

# ---------------------------------------------------------------------------
# Styling constants
# ---------------------------------------------------------------------------
TEAL = RGBColor(0, 128, 128)
TEAL_HEX = (0, 128, 128)
BLACK = RGBColor(0, 0, 0)


# ---------------------------------------------------------------------------
# Data loading
# ---------------------------------------------------------------------------
def load_experience_data(path: str) -> dict:
    """Load structured experience data from the given JSON file."""
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def _label(text: str) -> str:
    """Strip trailing colon from a bullet label."""
    return text.rstrip(": ").rstrip(":")


def format_resume_for_prompt(exp_data: dict) -> str:
    c = exp_data["contact"]
    lines = [
        c["name"],
        f"{c['address']} | {c['phone']} | {c['email']}",
        f"Headline: {exp_data['default_headline']}",
        f"\nSummary: {exp_data['default_summary']}",
        "\nWORK EXPERIENCE",
    ]
    for job in exp_data["experience"]:
        lines.append(f"\n{job['company']} -- {job['dates']}")
        for role in job["roles"]:
            lines.append(f"\n  {role['title']}")
            for b in role["bullets"]:
                tags = ", ".join(b.get("tags", []))
                lines.append(f"  - {_label(b['label'])}: {b['text']}  [tags: {tags}]")
    lines.append("\nSKILLS")
    for cat, skills in exp_data["skills"].items():
        lines.append(f"  {cat}: {', '.join(skills)}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# API calls
# ---------------------------------------------------------------------------
def call_claude(
    job_description: str,
    exp_data: dict,
    api_key: str,
    on_status: Callable[[str], None] = lambda m: None,
) -> dict:
    """Call Claude to tailor resume content. Returns structured JSON data."""
    client = Anthropic(api_key=api_key)

    resume_text = format_resume_for_prompt(exp_data)
    cover_letter_example = exp_data.get("cover_letter_example", "")

    system_prompt = textwrap.dedent("""\
        You are a professional resume and cover letter writer. You will be given a
        comprehensive pool of work experience, skills, and achievements along with a job
        description. Your job is to tailor a resume and write a cover letter for the
        specific role by selecting and reordering the most relevant content.

        RULES:
        - Do NOT fabricate experience, skills, or accomplishments. Only reword and
          reorder existing content to better match the job description.
        - The cover letter should be written in a confident, technical, first-person
          voice. Use the example cover letter as a style guide.
        - Keep bullet point labels (the bold prefix before each bullet).
        - Each bullet has tags indicating its relevance areas -- use these to help
          determine relevance to the job description.
        - Return ONLY valid JSON matching the schema below. No markdown fences.""")

    user_prompt = f"""FULL EXPERIENCE POOL:
{resume_text}

EXAMPLE COVER LETTER (use as style/tone reference only):
{cover_letter_example}

JOB DESCRIPTION:
{job_description}

Return a JSON object with this exact structure:
{{
  "company": "<extracted company name>",
  "job_title": "<extracted job title>",
  "headline": "<tailored professional headline>",
  "summary": "<tailored summary paragraph>",
  "experience": [
    {{
      "company": "<company name>",
      "dates": "<dates>",
      "roles": [
        {{
          "title": "<role title>",
          "bullets": [
            {{"label": "<bold label>", "text": "<bullet text>"}},
            ...
          ]
        }}
      ]
    }}
  ],
  "skills": {{
    "<category>": "<comma-separated skills, reordered to front-load relevant ones>",
    ...
  }},
  "cover_letter": "<full cover letter body text, from Dear Hiring Manager to sign-off>"
}}

For the experience section:
- Keep the company/roles structure as provided
- Reorder bullets within each role to put the most relevant ones first
- Reword bullets to emphasize skills/technologies mentioned in the job description
- Keep ALL bullets (do not remove any)
- Keep the label: text format for each bullet

For skills: keep all existing skills but reorder within each category to front-load
the ones most relevant to this job description.

For bullet labels: do NOT include a colon in the label value. The colon is added
automatically. For example use "Architectural Scale" not "Architectural Scale:".

For the cover letter:
- Address it to "Dear Hiring Manager,"
- Be concise: 2-3 short paragraphs plus a short bulleted list (3-4 items, one line each)
- Write in David's confident, technical voice (see example)
- End with a brief closing and "Best regards, David Kreuter"
- Keep the total length under 350 words"""

    on_status("Calling Claude API...")
    response = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=4096,
        system=system_prompt,
        messages=[{"role": "user", "content": user_prompt}],
    )

    raw = response.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1]
        if raw.endswith("```"):
            raw = raw[:-3]
    on_status("Claude response received.")
    return json.loads(raw)


def proofread_with_openai(
    data: dict,
    api_key: str,
    on_status: Callable[[str], None] = lambda m: None,
) -> dict:
    """Send tailored content through OpenAI for proofreading."""
    client = OpenAI(api_key=api_key)

    content = json.dumps({
        "headline": data["headline"],
        "summary": data["summary"],
        "experience": data["experience"],
        "cover_letter": data["cover_letter"],
    }, indent=2)

    on_status("Sending to OpenAI for language review...")
    response = client.chat.completions.create(
        model="o3",
        messages=[
            {
                "role": "system",
                "content": textwrap.dedent("""\
                    You are a professional proofreader and copy editor for resumes and
                    cover letters. Fix grammar, spelling, punctuation, awkward phrasing,
                    and tense consistency. Improve clarity and conciseness where possible.

                    RULES:
                    - Do NOT change the meaning, add new content, or remove content.
                    - Do NOT change technical terms, product names, or acronyms.
                    - Do NOT alter the JSON structure or keys.
                    - Keep bullet label values unchanged (they are category names).
                    - Preserve all newline characters (\\n) exactly as they appear.
                    - Return ONLY the corrected JSON with the same structure. No markdown fences."""),
            },
            {
                "role": "user",
                "content": f"Proofread and correct the language in this JSON:\n\n{content}",
            },
        ],
    )

    raw = response.choices[0].message.content.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1]
        if raw.endswith("```"):
            raw = raw[:-3]

    corrected = json.loads(raw)

    def _fix_newlines(text):
        if isinstance(text, str):
            text = text.replace("\\n", "\n")
        return text

    data["headline"] = _fix_newlines(corrected.get("headline", data["headline"]))
    data["summary"] = _fix_newlines(corrected.get("summary", data["summary"]))
    data["experience"] = corrected.get("experience", data["experience"])
    data["cover_letter"] = _fix_newlines(corrected.get("cover_letter", data["cover_letter"]))

    on_status("Proofreading complete.")
    return data


# ---------------------------------------------------------------------------
# TXT output
# ---------------------------------------------------------------------------
def write_resume_txt(data: dict, exp_data: dict, path: str):
    c = exp_data["contact"]
    lines = [
        c["name"],
        f"{c['address']}  |  {c['phone']}  |  {c['email']}",
        "",
        data["headline"],
        "",
        data["summary"],
        "",
        "WORK EXPERIENCE",
        "=" * 60,
    ]
    for job in data["experience"]:
        lines.append(f"\n{job['company']}    {job['dates']}")
        for role in job["roles"]:
            lines.append(role["title"])
            lines.append("-" * 40)
            for b in role["bullets"]:
                lines.append(f"  * {_label(b['label'])}: {b['text']}")
            lines.append("")

    lines.append("SKILLS")
    lines.append("=" * 60)
    for cat, skills in data["skills"].items():
        lines.append(f"  {cat}: {skills}")

    edu = exp_data["education"]
    lines.append(f"\nEDUCATION\n{'=' * 60}")
    lines.append(f"  {edu['degree']}    {edu['date']}")
    lines.append(f"  {edu['school']}")

    lines.append(f"\nCERTIFICATIONS\n{'=' * 60}")
    for cert in exp_data.get("certifications", []):
        lines.append(f"  {cert['name']}")
        lines.append(f"  {cert['issuer']}")

    lines.append(f"\nVOLUNTEERING & LEADERSHIP\n{'=' * 60}")
    for v in exp_data.get("volunteering", []):
        lines.append(f"  {v['org']} -- {v['dates']}")
        lines.append(f"  {v['role']}")

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def write_cover_letter_txt(data: dict, exp_data: dict, path: str):
    c = exp_data["contact"]
    lines = [
        c["name"],
        f"{c['address']}  |  {c['phone']}  |  {c['email']}",
        "",
        data["cover_letter"],
    ]
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


# ---------------------------------------------------------------------------
# DOCX output
# ---------------------------------------------------------------------------
def _set_run(run, size=10, bold=False, color=BLACK, font="Calibri"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font


def _add_section_header(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text)
    _set_run(run, size=11, bold=True, color=TEAL)
    pPr = p._p.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "008080")


def write_resume_docx(data: dict, exp_data: dict, path: str):
    c = exp_data["contact"]
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(c["name"])
    _set_run(run, size=20, bold=True, color=TEAL)

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{c['address']}  \u2022  {c['phone']}  \u2022  {c['email']}")
    _set_run(run, size=9)

    # Headline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data["headline"])
    _set_run(run, size=11, bold=True)

    # Summary
    p = doc.add_paragraph()
    run = p.add_run(data["summary"])
    _set_run(run, size=10)

    # Experience
    _add_section_header(doc, "WORK EXPERIENCE")
    for job in data["experience"]:
        for ri, role in enumerate(job["roles"]):
            if ri == 0:
                p = doc.add_paragraph()
                p.space_before = Pt(6)
                p.space_after = Pt(0)
                run = p.add_run(job["company"])
                _set_run(run, size=10, bold=True)
                run = p.add_run(f"    {job['dates']}")
                _set_run(run, size=9)

            p = doc.add_paragraph()
            p.space_before = Pt(0 if ri == 0 else 4)
            p.space_after = Pt(2)
            run = p.add_run(role["title"])
            _set_run(run, size=10, bold=True)

            for b in role["bullets"]:
                p = doc.add_paragraph(style="List Bullet")
                p.space_before = Pt(0)
                p.space_after = Pt(1)
                run = p.add_run(f"{_label(b['label'])}: ")
                _set_run(run, size=9, bold=True)
                run = p.add_run(b["text"])
                _set_run(run, size=9)

    # Skills
    _add_section_header(doc, "SKILLS")
    for cat, skills in data["skills"].items():
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(1)
        run = p.add_run(f"{cat}: ")
        _set_run(run, size=9, bold=True)
        run = p.add_run(skills)
        _set_run(run, size=9)

    # Education
    edu = exp_data["education"]
    _add_section_header(doc, "EDUCATION")
    p = doc.add_paragraph()
    p.space_before = Pt(0)
    p.space_after = Pt(0)
    run = p.add_run(edu["degree"])
    _set_run(run, size=10, bold=True)
    run = p.add_run(f"\t{edu['date']}")
    _set_run(run, size=9)
    p = doc.add_paragraph()
    run = p.add_run(edu["school"])
    _set_run(run, size=9)

    # Certifications
    _add_section_header(doc, "CERTIFICATIONS")
    for cert in exp_data.get("certifications", []):
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(0)
        run = p.add_run(cert["name"])
        _set_run(run, size=10, bold=True)
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(2)
        run = p.add_run(cert["issuer"])
        _set_run(run, size=9)

    # Volunteering
    _add_section_header(doc, "VOLUNTEERING & LEADERSHIP")
    for v in exp_data.get("volunteering", []):
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(0)
        run = p.add_run(v["org"])
        _set_run(run, size=10, bold=True)
        run = p.add_run(f"    {v['dates']}")
        _set_run(run, size=9)
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(2)
        run = p.add_run(v["role"])
        _set_run(run, size=9)

    doc.save(path)


def write_cover_letter_docx(data: dict, exp_data: dict, path: str):
    c = exp_data["contact"]
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(4)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(c["name"])
    _set_run(run, size=20, bold=True, color=TEAL)

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{c['address']}  \u2022  {c['phone']}  \u2022  {c['email']}")
    _set_run(run, size=9)

    # Horizontal rule
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "008080")

    # Cover letter body
    body = data["cover_letter"]
    paragraphs = body.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        bullet_lines = [l for l in para_text.split("\n") if l.strip().startswith("- ")]
        non_bullet_lines = [l for l in para_text.split("\n") if not l.strip().startswith("- ")]

        if non_bullet_lines:
            non_bullet_text = " ".join(l.strip() for l in non_bullet_lines if l.strip())
            if non_bullet_text:
                p = doc.add_paragraph()
                run = p.add_run(non_bullet_text)
                _set_run(run, size=9)

        for bl in bullet_lines:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(bl.strip().lstrip("- "))
            _set_run(run, size=9)

    doc.save(path)


# ---------------------------------------------------------------------------
# PDF output
# ---------------------------------------------------------------------------
def _pdf_safe(text: str) -> str:
    replacements = {
        "\u2022": "-",
        "\u2013": "-",
        "\u2014": "--",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2026": "...",
    }
    for char, repl in replacements.items():
        text = text.replace(char, repl)
    return text.encode("latin-1", errors="replace").decode("latin-1")


class ResumePDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)

    def section_header(self, text):
        self.set_font("Helvetica", "B", 11)
        self.set_text_color(*TEAL_HEX)
        self.cell(0, 7, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        self.set_draw_color(*TEAL_HEX)
        self.set_line_width(0.3)
        y = self.get_y()
        self.line(self.l_margin, y, self.w - self.r_margin, y)
        self.ln(3)
        self.set_text_color(0, 0, 0)

    def add_bullet(self, label, text):
        label = _pdf_safe(label).rstrip(":")
        text = _pdf_safe(text)
        lh = 4
        bullet_indent = 4
        indent_x = self.l_margin + bullet_indent
        wrap_w = self.w - self.r_margin - indent_x

        full_text = f"{label}: {text}"
        self.set_font("Helvetica", "", 8.5)
        words = full_text.split(" ")
        lines = []
        current = ""
        for w in words:
            test = f"{current} {w}".strip()
            if self.get_string_width(test) <= wrap_w:
                current = test
            else:
                if current:
                    lines.append(current)
                current = w
        if current:
            lines.append(current)

        self.set_font("Helvetica", "", 8.5)
        self.cell(bullet_indent, lh, "- ")

        bold_prefix = f"{label}: "
        bold_len = len(bold_prefix)
        chars_rendered = 0
        for i, line in enumerate(lines):
            if i > 0:
                self.set_x(indent_x)
            if chars_rendered < bold_len:
                bold_chars_left = bold_len - chars_rendered
                if bold_chars_left >= len(line):
                    self.set_font("Helvetica", "B", 8.5)
                    self.cell(wrap_w, lh, line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                else:
                    bold_part = line[:bold_chars_left]
                    normal_part = line[bold_chars_left:]
                    self.set_font("Helvetica", "B", 8.5)
                    self.cell(self.get_string_width(bold_part), lh, bold_part)
                    self.set_font("Helvetica", "", 8.5)
                    self.cell(wrap_w - self.get_string_width(bold_part), lh, normal_part,
                              new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            else:
                self.set_font("Helvetica", "", 8.5)
                self.cell(wrap_w, lh, line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            chars_rendered += len(line) + 1


def write_resume_pdf(data: dict, exp_data: dict, path: str):
    c = exp_data["contact"]
    pdf = ResumePDF()
    pdf.add_page()
    pdf.set_margins(18, 12, 18)

    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(*TEAL_HEX)
    pdf.cell(0, 10, c["name"], align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(0, 0, 0)
    contact = f"{c['address']}  |  {c['phone']}  |  {c['email']}"
    pdf.cell(0, 5, contact, align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 7, _pdf_safe(data["headline"]), align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    pdf.set_font("Helvetica", "", 9)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 4, _pdf_safe(data["summary"]), align="L")
    pdf.ln(3)

    pdf.section_header("WORK EXPERIENCE")
    for job in data["experience"]:
        for ri, role in enumerate(job["roles"]):
            if ri == 0:
                pdf.set_font("Helvetica", "B", 10)
                company = _pdf_safe(job["company"])
                dates = _pdf_safe(job["dates"])
                usable = pdf.w - pdf.l_margin - pdf.r_margin
                pdf.cell(usable / 2, 6, company)
                pdf.cell(usable / 2, 6, dates, align="R", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

            pdf.set_font("Helvetica", "B", 9.5)
            pdf.cell(0, 5, _pdf_safe(role["title"]), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(1)
            for b in role["bullets"]:
                pdf.add_bullet(b["label"], b["text"])
            if ri < len(job["roles"]) - 1:
                pdf.ln(2)
        pdf.ln(3)

    pdf.section_header("SKILLS")
    for cat, skills in data["skills"].items():
        full_text = f"{cat}: {_pdf_safe(skills)}"
        prefix = f"{cat}: "
        prefix_len = len(prefix)
        usable_w = pdf.w - pdf.l_margin - pdf.r_margin

        pdf.set_font("Helvetica", "", 9)
        words = full_text.split(" ")
        lines = []
        current = ""
        for w in words:
            test = f"{current} {w}".strip()
            if pdf.get_string_width(test) <= usable_w:
                current = test
            else:
                if current:
                    lines.append(current)
                current = w
        if current:
            lines.append(current)

        chars_rendered = 0
        for i, line in enumerate(lines):
            if i > 0:
                pdf.set_x(pdf.l_margin)
            if chars_rendered < prefix_len:
                bold_left = prefix_len - chars_rendered
                if bold_left >= len(line):
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.cell(usable_w, 4, line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                else:
                    bold_part = line[:bold_left]
                    normal_part = line[bold_left:]
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.cell(pdf.get_string_width(bold_part), 4, bold_part)
                    pdf.set_font("Helvetica", "", 9)
                    pdf.cell(usable_w - pdf.get_string_width(bold_part), 4, normal_part,
                             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            else:
                pdf.set_font("Helvetica", "", 9)
                pdf.cell(usable_w, 4, line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            chars_rendered += len(line) + 1
    pdf.ln(2)

    edu = exp_data["education"]
    pdf.section_header("EDUCATION")
    usable = pdf.w - pdf.l_margin - pdf.r_margin
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(usable / 2, 5, edu["degree"])
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(usable / 2, 5, edu["date"], align="R", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 5, edu["school"], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    pdf.section_header("CERTIFICATIONS")
    for cert in exp_data.get("certifications", []):
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 5, cert["name"], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(0, 4, cert["issuer"], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    pdf.section_header("VOLUNTEERING & LEADERSHIP")
    usable = pdf.w - pdf.l_margin - pdf.r_margin
    for v in exp_data.get("volunteering", []):
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(usable / 2, 4, v["org"])
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(usable / 2, 4, v["dates"], align="R", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", "", 8.5)
        pdf.cell(0, 4, v["role"], new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.output(path)


def write_cover_letter_pdf(data: dict, exp_data: dict, path: str):
    c = exp_data["contact"]
    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(20, 15, 20)
    pdf.set_auto_page_break(auto=True, margin=15)

    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(*TEAL_HEX)
    pdf.cell(0, 10, c["name"], align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(0, 0, 0)
    contact = f"{c['address']}  |  {c['phone']}  |  {c['email']}"
    pdf.cell(0, 5, contact, align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.ln(3)
    pdf.set_draw_color(*TEAL_HEX)
    pdf.set_line_width(0.5)
    y = pdf.get_y()
    pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
    pdf.ln(8)

    pdf.set_font("Helvetica", "", 9)
    body = _pdf_safe(data["cover_letter"])

    signoff_match = re.search(
        r'(Best regards|Sincerely|Kind regards|Regards|Warm regards),?\s*\n*\s*(David Kreuter)\s*$',
        body, re.IGNORECASE
    )
    if signoff_match:
        body = body[:signoff_match.start()].rstrip()
        signoff_line = signoff_match.group(1) + ","
        name_line = signoff_match.group(2)
    else:
        signoff_line = None
        name_line = None

    paragraphs = body.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        lines = para_text.split("\n")
        bullet_lines = [l for l in lines if l.strip().startswith("- ")]
        non_bullet_lines = [l for l in lines if not l.strip().startswith("- ")]

        if non_bullet_lines:
            text = " ".join(l.strip() for l in non_bullet_lines if l.strip())
            if text:
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 4.5, text)
                pdf.ln(2)

        for bl in bullet_lines:
            bl_text = bl.strip().lstrip("- ")
            indent = 5
            usable_w = pdf.w - pdf.l_margin - pdf.r_margin
            pdf.set_x(pdf.l_margin)
            pdf.cell(indent, 4.5, "- ")
            pdf.multi_cell(usable_w - indent, 4.5, bl_text)

    if signoff_line:
        pdf.ln(4)
        pdf.cell(0, 4.5, signoff_line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(2)
        pdf.cell(0, 4.5, name_line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.output(path)


# ---------------------------------------------------------------------------
# Orchestrator
# ---------------------------------------------------------------------------
def generate_application(
    job_description: str,
    exp_data: dict,
    options: GenerationOptions,
    anthropic_key: str,
    openai_key: str,
    output_dir: str,
    naming_template: str,
    conflict_mode: str,
    on_status: Callable[[str], None] = lambda m: None,
) -> tuple[dict, list[GeneratedFile]]:
    """Full generation pipeline: Claude -> OpenAI -> write files.

    Returns (tailored_data, list_of_generated_files).
    """
    # Step 1: Call Claude
    data = call_claude(job_description, exp_data, anthropic_key, on_status)

    # Step 2: Proofread with OpenAI
    data = proofread_with_openai(data, openai_key, on_status)

    # Step 3: Build file paths (late resolution — after we know job_title/company)
    company = data.get("company", "Company")
    job_title = data.get("job_title", "Role")
    resume_base, cover_base = build_base_paths(
        naming_template, job_title, company, output_dir
    )

    on_status(f"Generating files for: {job_title} at {company}")

    generated_files: list[GeneratedFile] = []

    # Step 4: Write selected files
    writers = []
    if options.generate_resume:
        if options.format_txt:
            writers.append(("resume", "txt", write_resume_txt, resume_base))
        if options.format_docx:
            writers.append(("resume", "docx", write_resume_docx, resume_base))
        if options.format_pdf:
            writers.append(("resume", "pdf", write_resume_pdf, resume_base))

    if options.generate_cover_letter:
        if options.format_txt:
            writers.append(("cover_letter", "txt", write_cover_letter_txt, cover_base))
        if options.format_docx:
            writers.append(("cover_letter", "docx", write_cover_letter_docx, cover_base))
        if options.format_pdf:
            writers.append(("cover_letter", "pdf", write_cover_letter_pdf, cover_base))

    for file_type, fmt, write_fn, base in writers:
        path = resolve_filename(base, f".{fmt}", conflict_mode)
        write_fn(data, exp_data, path)
        generated_files.append(GeneratedFile(
            file_type=file_type, format=fmt, path=path
        ))
        on_status(f"  [OK] {os.path.basename(path)}")

    on_status(f"Done! {len(generated_files)} files created.")
    return data, generated_files
