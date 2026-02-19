#!/usr/bin/env python3
"""
Job Application Tailoring Script
Paste a job description, get a tailored resume and cover letter in .txt, .docx, and .pdf.
Requires: ANTHROPIC_API_KEY and OPENAI_API_KEY environment variables
"""

import json
import os
import sys
import textwrap

from anthropic import Anthropic
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


# ──────────────────────────────────────────────────────────────────────────────
# LOAD EXPERIENCE DATA
# ──────────────────────────────────────────────────────────────────────────────

def load_experience_data() -> dict:
    """Load the structured experience data from experience_data.json."""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    path = os.path.join(base_dir, "experience_data.json")
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


# ──────────────────────────────────────────────────────────────────────────────
# STYLING CONSTANTS
# ──────────────────────────────────────────────────────────────────────────────

TEAL = RGBColor(0, 128, 128)
TEAL_HEX = (0, 128, 128)
BLACK = RGBColor(0, 0, 0)


# ──────────────────────────────────────────────────────────────────────────────
# CLAUDE API
# ──────────────────────────────────────────────────────────────────────────────

def call_claude(job_description: str, exp_data: dict) -> dict:
    """Send experience data + job description to Claude, return structured tailored content."""
    client = Anthropic()

    resume_text = format_resume_for_prompt(exp_data)
    cover_letter_example = exp_data.get("cover_letter_example", "")

    system_prompt = textwrap.dedent("""\
        You are a professional resume and cover letter writer. You will be given a
        comprehensive pool of work experience, skills, and achievements along with a job
        description. Your job is to tailor a resume and write a cover letter for the
        specific role by selecting and reordering the most relevant content.

        RULES:
        - Do NOT fabricate experience, skills, or accomplishments. Only reword and
          reorder existing content to better match the job description.
        - The cover letter should be written in a confident, technical, first-person
          voice. Use the example cover letter as a style guide.
        - Keep bullet point labels (the bold prefix before each bullet).
        - Each bullet has tags indicating its relevance areas -- use these to help
          determine relevance to the job description.
        - Return ONLY valid JSON matching the schema below. No markdown fences.""")

    user_prompt = f"""FULL EXPERIENCE POOL:
{resume_text}

EXAMPLE COVER LETTER (use as style/tone reference only):
{cover_letter_example}

JOB DESCRIPTION:
{job_description}

Return a JSON object with this exact structure:
{{
  "company": "<extracted company name>",
  "job_title": "<extracted job title>",
  "headline": "<tailored professional headline>",
  "summary": "<tailored summary paragraph>",
  "experience": [
    {{
      "company": "<company name>",
      "dates": "<dates>",
      "roles": [
        {{
          "title": "<role title>",
          "bullets": [
            {{"label": "<bold label>", "text": "<bullet text>"}},
            ...
          ]
        }}
      ]
    }}
  ],
  "skills": {{
    "<category>": "<comma-separated skills, reordered to front-load relevant ones>",
    ...
  }},
  "cover_letter": "<full cover letter body text, from Dear Hiring Manager to sign-off>"
}}

For the experience section:
- Keep the company/roles structure as provided
- Reorder bullets within each role to put the most relevant ones first
- Reword bullets to emphasize skills/technologies mentioned in the job description
- Keep ALL bullets (do not remove any)
- Keep the label: text format for each bullet

For skills: keep all existing skills but reorder within each category to front-load
the ones most relevant to this job description.

For bullet labels: do NOT include a colon in the label value. The colon is added
automatically. For example use "Architectural Scale" not "Architectural Scale:".

For the cover letter:
- Address it to "Dear Hiring Manager,"
- Be concise: 2-3 short paragraphs plus a short bulleted list (3-4 items, one line each)
- Write in David's confident, technical voice (see example)
- End with a brief closing and "Best regards, David Kreuter"
- Keep the total length under 350 words"""

    print("Calling Claude API...")
    response = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=4096,
        system=system_prompt,
        messages=[{"role": "user", "content": user_prompt}],
    )

    raw = response.content[0].text.strip()
    # Strip markdown fences if present
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1]
        if raw.endswith("```"):
            raw = raw[:-3]
    return json.loads(raw)


def proofread_with_openai(data: dict) -> dict:
    """Send the tailored resume and cover letter through OpenAI for language correction."""
    client = OpenAI()

    # Build a flat text representation of all content to proofread
    content = json.dumps({
        "headline": data["headline"],
        "summary": data["summary"],
        "experience": data["experience"],
        "cover_letter": data["cover_letter"],
    }, indent=2)

    print("Sending to OpenAI for language review...")
    response = client.chat.completions.create(
        model="o3",
        messages=[
            {
                "role": "system",
                "content": textwrap.dedent("""\
                    You are a professional proofreader and copy editor for resumes and
                    cover letters. Fix grammar, spelling, punctuation, awkward phrasing,
                    and tense consistency. Improve clarity and conciseness where possible.

                    RULES:
                    - Do NOT change the meaning, add new content, or remove content.
                    - Do NOT change technical terms, product names, or acronyms.
                    - Do NOT alter the JSON structure or keys.
                    - Keep bullet label values unchanged (they are category names).
                    - Preserve all newline characters (\\n) exactly as they appear.
                    - Return ONLY the corrected JSON with the same structure. No markdown fences."""),
            },
            {
                "role": "user",
                "content": f"Proofread and correct the language in this JSON:\n\n{content}",
            },
        ],
    )

    raw = response.choices[0].message.content.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1]
        if raw.endswith("```"):
            raw = raw[:-3]

    corrected = json.loads(raw)

    def _fix_newlines(text):
        """Fix double-escaped newlines that LLMs sometimes produce."""
        if isinstance(text, str):
            # Replace literal backslash-n with actual newline
            text = text.replace("\\n", "\n")
        return text

    # Merge corrected fields back into the original data
    data["headline"] = _fix_newlines(corrected.get("headline", data["headline"]))
    data["summary"] = _fix_newlines(corrected.get("summary", data["summary"]))
    data["experience"] = corrected.get("experience", data["experience"])
    data["cover_letter"] = _fix_newlines(corrected.get("cover_letter", data["cover_letter"]))

    return data


def _label(text: str) -> str:
    """Strip trailing colon from a bullet label to avoid double colons."""
    return text.rstrip(": ").rstrip(":")


def format_resume_for_prompt(exp_data: dict) -> str:
    """Format the experience data into readable text for the prompt."""
    c = exp_data["contact"]
    lines = [
        c["name"],
        f"{c['address']} | {c['phone']} | {c['email']}",
        f"Headline: {exp_data['default_headline']}",
        f"\nSummary: {exp_data['default_summary']}",
        "\nWORK EXPERIENCE",
    ]
    for job in exp_data["experience"]:
        lines.append(f"\n{job['company']} -- {job['dates']}")
        for role in job["roles"]:
            lines.append(f"\n  {role['title']}")
            for b in role["bullets"]:
                tags = ", ".join(b.get("tags", []))
                lines.append(f"  - {_label(b['label'])}: {b['text']}  [tags: {tags}]")

    lines.append("\nSKILLS")
    for cat, skills in exp_data["skills"].items():
        lines.append(f"  {cat}: {', '.join(skills)}")

    return "\n".join(lines)


# ──────────────────────────────────────────────────────────────────────────────
# TXT OUTPUT
# ──────────────────────────────────────────────────────────────────────────────

def write_resume_txt(data: dict, exp_data: dict, path: str):
    c = exp_data["contact"]
    lines = [
        c["name"],
        f"{c['address']}  |  {c['phone']}  |  {c['email']}",
        "",
        data["headline"],
        "",
        data["summary"],
        "",
        "WORK EXPERIENCE",
        "=" * 60,
    ]
    for job in data["experience"]:
        lines.append(f"\n{job['company']}    {job['dates']}")
        for role in job["roles"]:
            lines.append(role["title"])
            lines.append("-" * 40)
            for b in role["bullets"]:
                lines.append(f"  * {_label(b['label'])}: {b['text']}")
            lines.append("")

    lines.append(f"SKILLS")
    lines.append("=" * 60)
    for cat, skills in data["skills"].items():
        lines.append(f"  {cat}: {skills}")

    edu = exp_data["education"]
    lines.append(f"\nEDUCATION\n{'=' * 60}")
    lines.append(f"  {edu['degree']}    {edu['date']}")
    lines.append(f"  {edu['school']}")

    lines.append(f"\nCERTIFICATIONS\n{'=' * 60}")
    for cert in exp_data.get("certifications", []):
        lines.append(f"  {cert['name']}")
        lines.append(f"  {cert['issuer']}")

    lines.append(f"\nVOLUNTEERING & LEADERSHIP\n{'=' * 60}")
    for v in exp_data.get("volunteering", []):
        lines.append(f"  {v['org']} -- {v['dates']}")
        lines.append(f"  {v['role']}")

    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


def write_cover_letter_txt(data: dict, exp_data: dict, path: str):
    c = exp_data["contact"]
    lines = [
        c["name"],
        f"{c['address']}  |  {c['phone']}  |  {c['email']}",
        "",
        data["cover_letter"],
    ]
    with open(path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))


# ──────────────────────────────────────────────────────────────────────────────
# DOCX OUTPUT
# ──────────────────────────────────────────────────────────────────────────────

def _set_run(run, size=10, bold=False, color=BLACK, font="Calibri"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font


def _add_section_header(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text)
    _set_run(run, size=11, bold=True, color=TEAL)
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from lxml import etree
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "008080")


def write_resume_docx(data: dict, exp_data: dict, path: str):
    c = exp_data["contact"]
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(c["name"])
    _set_run(run, size=20, bold=True, color=TEAL)

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{c['address']}  \u2022  {c['phone']}  \u2022  {c['email']}")
    _set_run(run, size=9)

    # Headline
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data["headline"])
    _set_run(run, size=11, bold=True)

    # Summary
    p = doc.add_paragraph()
    run = p.add_run(data["summary"])
    _set_run(run, size=10)

    # Experience
    _add_section_header(doc, "WORK EXPERIENCE")
    for job in data["experience"]:
        for ri, role in enumerate(job["roles"]):
            # Company + dates on first role only
            if ri == 0:
                p = doc.add_paragraph()
                p.space_before = Pt(6)
                p.space_after = Pt(0)
                run = p.add_run(job["company"])
                _set_run(run, size=10, bold=True)
                run = p.add_run(f"    {job['dates']}")
                _set_run(run, size=9)

            # Title
            p = doc.add_paragraph()
            p.space_before = Pt(0 if ri == 0 else 4)
            p.space_after = Pt(2)
            run = p.add_run(role["title"])
            _set_run(run, size=10, bold=True)

            for b in role["bullets"]:
                p = doc.add_paragraph(style="List Bullet")
                p.space_before = Pt(0)
                p.space_after = Pt(1)
                run = p.add_run(f"{_label(b['label'])}: ")
                _set_run(run, size=9, bold=True)
                run = p.add_run(b["text"])
                _set_run(run, size=9)

    # Skills
    _add_section_header(doc, "SKILLS")
    for cat, skills in data["skills"].items():
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(1)
        run = p.add_run(f"{cat}: ")
        _set_run(run, size=9, bold=True)
        run = p.add_run(skills)
        _set_run(run, size=9)

    # Education
    edu = exp_data["education"]
    _add_section_header(doc, "EDUCATION")
    p = doc.add_paragraph()
    p.space_before = Pt(0)
    p.space_after = Pt(0)
    run = p.add_run(edu["degree"])
    _set_run(run, size=10, bold=True)
    # Right-align date using a tab stop
    run = p.add_run(f"\t{edu['date']}")
    _set_run(run, size=9)
    p = doc.add_paragraph()
    run = p.add_run(edu["school"])
    _set_run(run, size=9)

    # Certifications
    _add_section_header(doc, "CERTIFICATIONS")
    for cert in exp_data.get("certifications", []):
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(0)
        run = p.add_run(cert["name"])
        _set_run(run, size=10, bold=True)
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(2)
        run = p.add_run(cert["issuer"])
        _set_run(run, size=9)

    # Volunteering
    _add_section_header(doc, "VOLUNTEERING & LEADERSHIP")
    for v in exp_data.get("volunteering", []):
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(0)
        run = p.add_run(v["org"])
        _set_run(run, size=10, bold=True)
        run = p.add_run(f"    {v['dates']}")
        _set_run(run, size=9)
        p = doc.add_paragraph()
        p.space_before = Pt(0)
        p.space_after = Pt(2)
        run = p.add_run(v["role"])
        _set_run(run, size=9)

    doc.save(path)


def write_cover_letter_docx(data: dict, exp_data: dict, path: str):
    c = exp_data["contact"]
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(4)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(c["name"])
    _set_run(run, size=20, bold=True, color=TEAL)

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{c['address']}  \u2022  {c['phone']}  \u2022  {c['email']}")
    _set_run(run, size=9)

    # Horizontal rule
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(12)
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from lxml import etree
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "008080")

    # Cover letter body
    body = data["cover_letter"]
    paragraphs = body.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        bullet_lines = [l for l in para_text.split("\n") if l.strip().startswith("- ")]
        non_bullet_lines = [l for l in para_text.split("\n") if not l.strip().startswith("- ")]

        if non_bullet_lines:
            non_bullet_text = " ".join(l.strip() for l in non_bullet_lines if l.strip())
            if non_bullet_text:
                p = doc.add_paragraph()
                run = p.add_run(non_bullet_text)
                _set_run(run, size=9)

        for bl in bullet_lines:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(bl.strip().lstrip("- "))
            _set_run(run, size=9)

    doc.save(path)


# ──────────────────────────────────────────────────────────────────────────────
# PDF OUTPUT
# ──────────────────────────────────────────────────────────────────────────────

def _pdf_safe(text: str) -> str:
    """Replace Unicode characters that latin-1 can't encode."""
    replacements = {
        "\u2022": "-",   # bullet
        "\u2013": "-",   # en-dash
        "\u2014": "--",  # em-dash
        "\u2018": "'",   # left single quote
        "\u2019": "'",   # right single quote
        "\u201c": '"',   # left double quote
        "\u201d": '"',   # right double quote
        "\u2026": "...", # ellipsis
    }
    for char, repl in replacements.items():
        text = text.replace(char, repl)
    return text.encode("latin-1", errors="replace").decode("latin-1")


class ResumePDF(FPDF):
    def __init__(self):
        super().__init__()
        self.set_auto_page_break(auto=True, margin=15)

    def section_header(self, text):
        self.set_font("Helvetica", "B", 11)
        self.set_text_color(*TEAL_HEX)
        self.cell(0, 7, text, ln=1)
        self.set_draw_color(*TEAL_HEX)
        self.set_line_width(0.3)
        y = self.get_y()
        self.line(self.l_margin, y, self.w - self.r_margin, y)
        self.ln(3)
        self.set_text_color(0, 0, 0)

    def add_bullet(self, label, text):
        label = _pdf_safe(label).rstrip(":")
        text = _pdf_safe(text)
        lh = 4
        bullet_indent = 4
        indent_x = self.l_margin + bullet_indent
        wrap_w = self.w - self.r_margin - indent_x

        full_text = f"{label}: {text}"

        self.set_font("Helvetica", "", 8.5)
        words = full_text.split(" ")
        lines = []
        current = ""
        for w in words:
            test = f"{current} {w}".strip()
            if self.get_string_width(test) <= wrap_w:
                current = test
            else:
                if current:
                    lines.append(current)
                current = w
        if current:
            lines.append(current)

        # Render "- " prefix
        self.set_font("Helvetica", "", 8.5)
        self.cell(bullet_indent, lh, "- ")

        # Render each wrapped line at indent_x
        bold_prefix = f"{label}: "
        bold_len = len(bold_prefix)
        chars_rendered = 0
        for i, line in enumerate(lines):
            if i > 0:
                self.set_x(indent_x)
            if chars_rendered < bold_len:
                bold_chars_left = bold_len - chars_rendered
                if bold_chars_left >= len(line):
                    self.set_font("Helvetica", "B", 8.5)
                    self.cell(wrap_w, lh, line, ln=1)
                else:
                    bold_part = line[:bold_chars_left]
                    normal_part = line[bold_chars_left:]
                    self.set_font("Helvetica", "B", 8.5)
                    self.cell(self.get_string_width(bold_part), lh, bold_part)
                    self.set_font("Helvetica", "", 8.5)
                    self.cell(wrap_w - self.get_string_width(bold_part), lh, normal_part, ln=1)
            else:
                self.set_font("Helvetica", "", 8.5)
                self.cell(wrap_w, lh, line, ln=1)
            chars_rendered += len(line) + 1


def write_resume_pdf(data: dict, exp_data: dict, path: str):
    c = exp_data["contact"]
    pdf = ResumePDF()
    pdf.add_page()
    pdf.set_margins(18, 12, 18)

    # Name
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(*TEAL_HEX)
    pdf.cell(0, 10, c["name"], ln=1, align="C")

    # Contact
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(0, 0, 0)
    contact = f"{c['address']}  |  {c['phone']}  |  {c['email']}"
    pdf.cell(0, 5, contact, ln=1, align="C")

    # Headline
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 7, _pdf_safe(data["headline"]), ln=1, align="C")
    pdf.ln(2)

    # Summary
    pdf.set_font("Helvetica", "", 9)
    pdf.multi_cell(0, 4, _pdf_safe(data["summary"]), align="L")
    pdf.ln(3)

    # Experience
    pdf.section_header("WORK EXPERIENCE")
    for job in data["experience"]:
        for ri, role in enumerate(job["roles"]):
            # Company + dates on first role only
            if ri == 0:
                pdf.set_font("Helvetica", "B", 10)
                company = _pdf_safe(job["company"])
                dates = _pdf_safe(job["dates"])
                usable = pdf.w - pdf.l_margin - pdf.r_margin
                pdf.cell(usable / 2, 6, company)
                pdf.cell(usable / 2, 6, dates, ln=1, align="R")

            pdf.set_font("Helvetica", "B", 9.5)
            pdf.cell(0, 5, _pdf_safe(role["title"]), ln=1)
            pdf.ln(1)
            for b in role["bullets"]:
                pdf.add_bullet(b["label"], b["text"])
            if ri < len(job["roles"]) - 1:
                pdf.ln(2)
        pdf.ln(3)

    # Skills
    pdf.section_header("SKILLS")
    for cat, skills in data["skills"].items():
        # Render "Category: skills..." with wrapped lines aligned to left margin
        full_text = f"{cat}: {_pdf_safe(skills)}"
        prefix = f"{cat}: "
        prefix_len = len(prefix)
        usable_w = pdf.w - pdf.l_margin - pdf.r_margin

        pdf.set_font("Helvetica", "", 9)
        words = full_text.split(" ")
        lines = []
        current = ""
        for w in words:
            test = f"{current} {w}".strip()
            if pdf.get_string_width(test) <= usable_w:
                current = test
            else:
                if current:
                    lines.append(current)
                current = w
        if current:
            lines.append(current)

        chars_rendered = 0
        for i, line in enumerate(lines):
            if i > 0:
                pdf.set_x(pdf.l_margin)
            if chars_rendered < prefix_len:
                bold_left = prefix_len - chars_rendered
                if bold_left >= len(line):
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.cell(usable_w, 4, line, ln=1)
                else:
                    bold_part = line[:bold_left]
                    normal_part = line[bold_left:]
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.cell(pdf.get_string_width(bold_part), 4, bold_part)
                    pdf.set_font("Helvetica", "", 9)
                    pdf.cell(usable_w - pdf.get_string_width(bold_part), 4, normal_part, ln=1)
            else:
                pdf.set_font("Helvetica", "", 9)
                pdf.cell(usable_w, 4, line, ln=1)
            chars_rendered += len(line) + 1
    pdf.ln(2)

    # Education
    edu = exp_data["education"]
    pdf.section_header("EDUCATION")
    usable = pdf.w - pdf.l_margin - pdf.r_margin
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(usable / 2, 5, edu["degree"])
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(usable / 2, 5, edu["date"], ln=1, align="R")
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 5, edu["school"], ln=1)
    pdf.ln(2)

    # Certifications
    pdf.section_header("CERTIFICATIONS")
    for cert in exp_data.get("certifications", []):
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 5, cert["name"], ln=1)
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(0, 4, cert["issuer"], ln=1)
    pdf.ln(2)

    # Volunteering
    pdf.section_header("VOLUNTEERING & LEADERSHIP")
    usable = pdf.w - pdf.l_margin - pdf.r_margin
    for v in exp_data.get("volunteering", []):
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(usable / 2, 4, v["org"])
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(usable / 2, 4, v["dates"], ln=1, align="R")
        pdf.set_font("Helvetica", "", 8.5)
        pdf.cell(0, 4, v["role"], ln=1)

    pdf.output(path)


def write_cover_letter_pdf(data: dict, exp_data: dict, path: str):
    c = exp_data["contact"]
    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(20, 15, 20)
    pdf.set_auto_page_break(auto=True, margin=15)

    # Name
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(*TEAL_HEX)
    pdf.cell(0, 10, c["name"], ln=1, align="C")

    # Contact
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(0, 0, 0)
    contact = f"{c['address']}  |  {c['phone']}  |  {c['email']}"
    pdf.cell(0, 5, contact, ln=1, align="C")

    # Teal rule
    pdf.ln(3)
    pdf.set_draw_color(*TEAL_HEX)
    pdf.set_line_width(0.5)
    y = pdf.get_y()
    pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
    pdf.ln(8)

    # Body - match contact info font style
    pdf.set_font("Helvetica", "", 9)
    body = _pdf_safe(data["cover_letter"])

    # Strip sign-off from body so we can render name on its own line
    import re
    signoff_match = re.search(
        r'(Best regards|Sincerely|Kind regards|Regards|Warm regards),?\s*\n*\s*(David Kreuter)\s*$',
        body, re.IGNORECASE
    )
    if signoff_match:
        body = body[:signoff_match.start()].rstrip()
        signoff_line = signoff_match.group(1) + ","
        name_line = signoff_match.group(2)
    else:
        signoff_line = None
        name_line = None

    paragraphs = body.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        lines = para_text.split("\n")
        bullet_lines = [l for l in lines if l.strip().startswith("- ")]
        non_bullet_lines = [l for l in lines if not l.strip().startswith("- ")]

        if non_bullet_lines:
            text = " ".join(l.strip() for l in non_bullet_lines if l.strip())
            if text:
                pdf.multi_cell(0, 4.5, text)
                pdf.ln(2)

        for bl in bullet_lines:
            bl_text = bl.strip().lstrip("- ")
            pdf.cell(5, 4.5, "- ")
            remaining = pdf.w - pdf.r_margin - pdf.get_x()
            pdf.multi_cell(remaining, 4.5, bl_text)

    # Render sign-off with name on its own line
    if signoff_line:
        pdf.ln(4)
        pdf.cell(0, 4.5, signoff_line, ln=1)
        pdf.ln(2)
        pdf.cell(0, 4.5, name_line, ln=1)

    pdf.output(path)


# ──────────────────────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────────────────────

def main():
    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("Error: ANTHROPIC_API_KEY environment variable is not set.")
        sys.exit(1)
    if not os.environ.get("OPENAI_API_KEY"):
        print("Error: OPENAI_API_KEY environment variable is not set.")
        sys.exit(1)

    exp_data = load_experience_data()
    print(f"Loaded experience data: {len(exp_data['experience'])} companies, "
          f"{sum(len(r['bullets']) for j in exp_data['experience'] for r in j['roles'])} bullets")

    print("\nPaste the job description below.")
    print("When done, press Ctrl+Z then Enter (Windows) or Ctrl+D (Mac/Linux):\n")

    lines = []
    try:
        while True:
            line = input()
            lines.append(line)
    except EOFError:
        pass

    job_description = "\n".join(lines).strip()
    if not job_description:
        print("No job description provided. Exiting.")
        sys.exit(1)

    print(f"\nReceived {len(job_description)} characters of job description.")

    data = call_claude(job_description, exp_data)
    data = proofread_with_openai(data)

    company = data.get("company", "Company")
    job_title = data.get("job_title", "Role")
    safe = lambda s: "".join(c for c in s if c.isalnum() or c in " -_&").strip()
    company = safe(company)
    job_title = safe(job_title)

    base_dir = os.path.dirname(os.path.abspath(__file__))
    resume_base = os.path.join(base_dir, f"{job_title} at {company} Resume")
    cover_base = os.path.join(base_dir, f"{job_title} at {company} Cover Letter")

    print(f"\nGenerating files for: {job_title} at {company}")

    # TXT
    write_resume_txt(data, exp_data, resume_base + ".txt")
    write_cover_letter_txt(data, exp_data, cover_base + ".txt")
    print("  [OK] .txt files")

    # DOCX
    write_resume_docx(data, exp_data, resume_base + ".docx")
    write_cover_letter_docx(data, exp_data, cover_base + ".docx")
    print("  [OK] .docx files")

    # PDF
    write_resume_pdf(data, exp_data, resume_base + ".pdf")
    write_cover_letter_pdf(data, exp_data, cover_base + ".pdf")
    print("  [OK] .pdf files")

    print(f"\nDone! 6 files created in {base_dir}")


if __name__ == "__main__":
    main()
